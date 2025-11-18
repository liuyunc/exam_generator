# docx_utils.py

import re

from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn
from typing import Dict, List, Sequence


def _normalize_options(raw: object) -> List[str]:
    """将 options 字段转换为规范的字符串列表，便于文档渲染。"""

    if raw is None:
        return []
    if isinstance(raw, str):
        return [line.strip() for line in raw.splitlines() if line.strip()]
    if isinstance(raw, Sequence):
        result = []
        for item in raw:
            text = str(item).strip()
            if text:
                result.append(text)
        return result
    return []


def _sanitize_math_markdown(text: str) -> str:
    """尽量将 Markdown/LaTeX 形式的公式转为可读文本，避免 Word 中保留 $ 符号。"""

    if not text:
        return ""

    def _clean_math_content(content: str) -> str:
        # 去除常见的 LaTeX 包裹符号
        content = re.sub(r"\\mathrm\{([^}]+)\}", r"\1", content)
        content = re.sub(r"\\operatorname\{([^}]+)\}", r"\1", content)
        content = re.sub(r"_\{([^}]+)\}", r"_\1", content)
        content = re.sub(r"\^\{([^}]+)\}", r"^\1", content)
        content = re.sub(r"\\text\{([^}]+)\}", r"\1", content)
        # 清理 \left / \right 之类的定界符标记
        content = re.sub(r"\\(left|right|bigl|bigr|Bigl|Bigr|biggl|biggr|Biggl|Biggr)", "", content)
        # 去掉剩余的反斜杠标记（如 \mathrm、\alpha -> alpha）
        content = re.sub(r"\\([a-zA-Z]+)", r"\1", content)
        return content.replace("{", "").replace("}", "")

    # 处理 $...$、$$...$$、\(...\)、\[...\] 包裹的公式
    inline_patterns = [
        r"\$(.+?)\$",
        r"\$\$(.+?)\$\$",
        r"\\\((.+?)\\\)",
        r"\\\[(.+?)\\\]",
    ]
    for pattern in inline_patterns:
        text = re.sub(pattern, lambda m: _clean_math_content(m.group(1)), text)
    # 再兜底清理未成对的 \mathrm{} 等
    text = _clean_math_content(text)
    return text


QUESTION_TYPE_MAP: Dict[str, str] = {
    "single_choice": "【单选题】",
    "multiple_choice": "【多选题】",
    "true_false": "【判断题】",
    "short_answer": "【简答题】",
}


def _render_question_type(raw: str) -> str:
    """兼容 minerU 解析与 GA 输出的题型标记，统一为中文展示。"""

    key = (raw or "").strip().lower()
    # 题面中可能出现 [single_choice] 形式，直接替换。
    bracket_clean = re.sub(r"^\[(.+)\]$", r"\1", key)
    mapped = QUESTION_TYPE_MAP.get(bracket_clean, "")
    return mapped or raw


def _replace_type_tokens(text: str) -> str:
    """将题干中的 [single_choice] 等占位替换为中文题型标签。"""

    if not text:
        return ""

    def _sub(match):
        inner = match.group(1).lower()
        return QUESTION_TYPE_MAP.get(inner, match.group(0))

    return re.sub(r"\[([a-z_]+)\]", _sub, text)


def sort_ga_pairs_by_type(ga_pairs: List[dict]) -> List[dict]:
    """按题型分类排序，顺序为单选、多选、判断、简答。"""

    order_map = {
        "single_choice": 0,
        "multiple_choice": 1,
        "true_false": 2,
        "short_answer": 3,
    }

    def sort_key(item_with_idx):
        idx, qa = item_with_idx
        q_type = (qa.get("question_type") or "").strip()
        return order_map.get(q_type, len(order_map)), idx

    return [qa for _, qa in sorted(enumerate(ga_pairs), key=sort_key)]


def build_docx_from_ga(ga_pairs, title: str = "培训考试题（含答案与原文引用）"):
    """
    ga_pairs: List[dict]，字段：
      id, question, question_type, options, ga_answer, difficulty,
      source_excerpt, source_locator, comment
    返回：python-docx 的 Document 对象
    """
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "宋体"
    style.element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    style.font.size = Pt(11)

    for heading_style in ["Heading 1", "Heading 2", "Heading 3"]:
        if heading_style in doc.styles:
            h_style = doc.styles[heading_style]
            h_style.font.name = "宋体"
            h_style.element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")

    # 标题
    doc.add_heading(title, level=1)

    # 一、试题（不含答案）
    doc.add_heading("一、试题（不含答案）", level=2)
    for idx, qa in enumerate(ga_pairs, start=1):
        q = _replace_type_tokens(
            _sanitize_math_markdown((qa.get("question") or "").strip())
        )
        question_type = _render_question_type(
            _sanitize_math_markdown((qa.get("question_type") or "").strip())
        )
        options = [_sanitize_math_markdown(opt) for opt in _normalize_options(qa.get("options"))]

        p = doc.add_paragraph()
        prefix = f"{idx}. "
        if question_type:
            prefix += f"{question_type} "
        p.add_run(f"{prefix}{q}")

        for opt in options:
            opt_p = doc.add_paragraph()
            opt_p.add_run(f"   {opt}")

    doc.add_page_break()

    # 二、参考答案与原文引用
    doc.add_heading("二、参考答案与原文引用", level=2)
    for idx, qa in enumerate(ga_pairs, start=1):
        q = _replace_type_tokens(
            _sanitize_math_markdown((qa.get("question") or "").strip())
        )
        a = _sanitize_math_markdown((qa.get("ga_answer") or "").strip())
        question_type = _render_question_type(
            _sanitize_math_markdown((qa.get("question_type") or "").strip())
        )
        options = [_sanitize_math_markdown(opt) for opt in _normalize_options(qa.get("options"))]
        difficulty = _sanitize_math_markdown((qa.get("difficulty") or "").strip())
        source_excerpt = _sanitize_math_markdown((qa.get("source_excerpt") or "").strip())
        locator = _sanitize_math_markdown((qa.get("source_locator") or "").strip())
        comment = _sanitize_math_markdown((qa.get("comment") or "").strip())

        doc.add_paragraph(f"{idx}. 题目：{q}")
        if question_type:
            doc.add_paragraph(f"   【题型】{question_type}")
        if options:
            doc.add_paragraph("   【选项】")
            for opt in options:
                doc.add_paragraph(f"      {opt}")
        doc.add_paragraph(f"   【参考答案】{a}")
        if difficulty:
            doc.add_paragraph(f"   【难度】{difficulty}")
        if locator:
            doc.add_paragraph(f"   【来源定位】{locator}")
        if source_excerpt:
            doc.add_paragraph(f"   【原文摘录】{source_excerpt}")
        if comment:
            doc.add_paragraph(f"   【命题说明】{comment}")
        doc.add_paragraph("")  # 空行分隔

    return doc
